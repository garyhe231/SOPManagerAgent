"""
Extract text content from PDF, DOCX, XLSX, and video files.
"""
import os
import tempfile
from pathlib import Path
from typing import Optional

import pdfplumber
from docx import Document
from openpyxl import load_workbook


def extract_pdf(filepath: str) -> str:
    lines = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text:
                lines.append(f"--- Page {i} ---\n{text.strip()}")
    return "\n\n".join(lines)


def extract_docx(filepath: str) -> str:
    doc = Document(filepath)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # Also capture table content
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_xlsx(filepath: str) -> str:
    wb = load_workbook(filepath, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            row_text = " | ".join(cells).strip(" |")
            if row_text.replace("|", "").strip():
                parts.append(row_text)
    return "\n".join(parts)


def extract_video(filepath: str) -> str:
    """Transcribe video/audio using Whisper."""
    try:
        import whisper
        model = whisper.load_model("base")
        result = model.transcribe(filepath)
        return result.get("text", "").strip()
    except Exception as e:
        return f"[Video transcription failed: {e}]"


def extract_text(filepath: str, filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return extract_docx(filepath)
    elif ext in (".xlsx", ".xls"):
        return extract_xlsx(filepath)
    elif ext in (".mp4", ".mov", ".avi", ".mkv", ".webm", ".m4v", ".mp3", ".wav", ".m4a"):
        return extract_video(filepath)
    else:
        # Try reading as plain text
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            return f"[Unable to extract text: {e}]"
